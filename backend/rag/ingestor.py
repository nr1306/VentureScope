"""
Document ingestion pipeline:
  raw bytes → parse text → chunk → embed → store in pgvector
"""
from __future__ import annotations
import asyncio
import io
import uuid

from sqlalchemy.ext.asyncio import AsyncSession

from config import settings
from db.models import Document, DocumentChunk
from db.session import commit_with_retry, flush_with_retry
from guardrails.input_guardrails import validate_company_input, validate_upload
from rag.chunker import split_text
from rag.embedder import EmbeddingError, embed_texts


SUPPORTED_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
}


class DocumentIngestionError(RuntimeError):
    """Raised when a document cannot be safely ingested."""


def _parse_document(filename: str, content_type: str, contents: bytes) -> str:
    """Extract plain text from uploaded file."""
    if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(contents))
        return "\n\n".join(
            page.extract_text() or "" for page in reader.pages
        ).strip()

    if (
        content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or filename.lower().endswith(".docx")
    ):
        import docx
        doc = docx.Document(io.BytesIO(contents))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # Plain text / markdown
    return contents.decode("utf-8", errors="replace")


def _dedupe_chunks(chunks: list[str]) -> list[str]:
    seen: set[str] = set()
    deduped: list[str] = []
    for chunk in chunks:
        normalized = " ".join(chunk.split()).strip().lower()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(chunk)
    return deduped


async def ingest_document(
    db: AsyncSession,
    company_name: str,
    filename: str,
    content_type: str,
    contents: bytes,
) -> tuple[str, int]:
    """
    Parse, chunk, embed, and store a document.
    Returns (document_id, chunk_count).
    """
    safe_company_name = company_name.strip()
    validate_company_input(safe_company_name)
    safe_filename = validate_upload(
        filename=filename,
        content_type=content_type,
        size_bytes=len(contents),
        contents=contents,
    )

    try:
        text = await asyncio.to_thread(_parse_document, safe_filename, content_type, contents)
    except Exception as exc:
        raise DocumentIngestionError(f"Failed to parse uploaded document '{safe_filename}'.") from exc

    if not text.strip():
        raise DocumentIngestionError("Uploaded document does not contain extractable text.")

    chunks = _dedupe_chunks(split_text(text))

    if len(chunks) > settings.max_chunks_per_doc:
        chunks = chunks[:settings.max_chunks_per_doc]
    if not chunks:
        raise DocumentIngestionError("No chunks were produced from the uploaded document.")

    doc_id = str(uuid.uuid4())
    doc = Document(
        id=doc_id,
        company_name=safe_company_name,
        filename=safe_filename,
        content_type=content_type,
        chunk_count=len(chunks),
    )

    try:
        db.add(doc)
        await flush_with_retry(db)

        # Embed in batches of 100
        batch_size = 100
        all_chunk_objs = []
        for batch_start in range(0, len(chunks), batch_size):
            batch = chunks[batch_start: batch_start + batch_size]
            try:
                embeddings = await embed_texts(batch)
            except EmbeddingError as exc:
                raise DocumentIngestionError(
                    f"Embedding failed while ingesting '{safe_filename}'."
                ) from exc

            if len(embeddings) != len(batch):
                raise DocumentIngestionError(
                    f"Embedding response size mismatch while ingesting '{safe_filename}'."
                )

            for i, (chunk_text, embedding) in enumerate(zip(batch, embeddings)):
                all_chunk_objs.append(
                    DocumentChunk(
                        id=str(uuid.uuid4()),
                        document_id=doc_id,
                        company_name=safe_company_name,
                        chunk_index=batch_start + i,
                        content=chunk_text,
                        embedding=embedding,
                        metadata_={"filename": safe_filename, "chunk_index": batch_start + i},
                    )
                )

        if not all_chunk_objs:
            raise DocumentIngestionError("No document chunks were generated for storage.")

        db.add_all(all_chunk_objs)
        await commit_with_retry(db)
    except Exception:
        await db.rollback()
        raise

    return doc_id, len(chunks)
